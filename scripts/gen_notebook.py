import json

_c = [0]

def md(text):
    _c[0] += 1
    t = text.strip('\n')
    lines = t.split('\n')
    src = [l + '\n' for l in lines[:-1]] + [lines[-1]]
    return {"cell_type": "markdown", "id": f"m{_c[0]:03d}", "metadata": {}, "source": src}

def code(text):
    _c[0] += 1
    t = text.strip('\n')
    lines = t.split('\n')
    src = [l + '\n' for l in lines[:-1]] + [lines[-1]]
    return {"cell_type": "code", "execution_count": None, "id": f"c{_c[0]:03d}", "metadata": {}, "outputs": [], "source": src}

cells = []

# ── INTRO ──────────────────────────────────────────────────────────────────
cells.append(md('''# Học Lập Trình Ứng Dụng Văn Phòng với Claude Code
### 3 Buổi Đầu Tiên — Dành cho người không chuyên lập trình

---

## Trước khi bắt đầu — Hiểu đúng vai trò của bạn

Bạn **không học để trở thành lập trình viên.**

Bạn học để trở thành người **biết chỉ đạo Claude** làm việc thay mình.

| | Lập trình viên | Bạn (với Claude) |
|---|---|---|
| Mục tiêu | Tự viết code từ đầu | Hiểu đủ để điều khiển Claude |
| Vai trò | Người thợ | Người chỉ đạo |
| Cần biết | Viết thành thạo | Đọc hiểu + mô tả đúng |
| Thời gian | Năm | Tuần |

Kiến thức lập trình bạn cần chỉ đủ để:
- Hiểu code Claude viết ra đang làm gì (đại ý, không cần từng dòng)
- Biết sửa những chỗ nhỏ (tên file, cột dữ liệu, con số)
- Đọc thông báo lỗi và mô tả lại cho Claude
- Biết cái gì có thể tự động hóa được

---

**Cài đặt một lần trước buổi 1:**
1. Cài Python: `python.org/downloads` → chọn bản mới nhất → cài mặc định
2. Cài VS Code: `code.visualstudio.com`
3. Mở VS Code → cài extension **Python** (tìm trong tab Extensions)'''))

# ── BUỔI 1 ─────────────────────────────────────────────────────────────────
cells.append(md('''---
---
# BUỔI 1 — Hiểu Cách Máy Tính Làm Việc và Chạy Code Đầu Tiên

**Mục tiêu buổi này:**
- Hiểu tại sao cần biết lập trình để dùng Claude hiệu quả
- Chạy được code Python đầu tiên
- Hiểu biến và kiểu dữ liệu khi đọc code Claude viết'''))

cells.append(md('''## Phần 1 — Tại sao cần hiểu code dù có Claude?

Claude giỏi viết code, nhưng Claude không biết:
- File Excel của bạn tên gì, để ở đâu
- Cột nào là doanh thu, cột nào là tên khách
- Kết quả "đúng" trông như thế nào với bạn

Nếu bạn mù hoàn toàn về code, bạn sẽ không biết Claude đang làm đúng hay sai.
Giống như thuê thợ sửa nhà mà không biết gì về xây dựng — bạn không thể kiểm tra chất lượng.

Bạn chỉ cần hiểu **vừa đủ** để làm người giám sát thông minh.'''))

cells.append(md('''## Phần 2 — Máy tính làm việc như thế nào?

Máy tính thực hiện lệnh **theo thứ tự từ trên xuống, chính xác tuyệt đối, không suy luận.**

Không có bước nào tự động xảy ra nếu bạn không viết ra.
Đây là lý do bạn phải mô tả yêu cầu cho Claude **thật cụ thể** — Claude cũng cần biết từng bước.'''))

cells.append(code('''# Máy tính chạy từng dòng, theo thứ tự, từ trên xuống
print("Bước 1: Mở file")
print("Bước 2: Đọc dữ liệu")
print("Bước 3: Tính toán")
print("Bước 4: Lưu kết quả")'''))

cells.append(md('''## Phần 3 — Chạy code đầu tiên

Nhấn nút ▶ ở góc trái ô code bên dưới để chạy.
Kết quả xuất hiện ngay bên dưới ô code.'''))

cells.append(code('''print("Xin chào!")
print("Tôi đang học dùng Claude cho công việc văn phòng")
print(2024 + 1)'''))

cells.append(md('''## Phần 4 — Biến: cách code lưu thông tin

Khi Claude viết code xử lý file Excel, code sẽ có các dòng như thế này:

```python
ten_file = "bao_cao_thang_5.xlsx"
thu_muc = "C:/Users/NhanVien/Desktop/BaoCao"
cot_doanh_thu = "D"
```

Đây là **biến** — hộp chứa thông tin, có đặt tên.

**Cú pháp:** `tên_biến = giá_trị`

Dấu `=` trong lập trình **không phải** "bằng nhau" — nó có nghĩa là **"gán giá trị vào"**.

**Tại sao bạn cần biết?**
Khi Claude viết code cho bạn, thường có vài dòng biến ở đầu file mà **bạn cần tự sửa** cho đúng với máy tính của mình. Bạn chỉ cần sửa phần trong dấu ngoặc kép.

**Quy tắc đặt tên biến:**
- Dùng chữ thường, nối bằng dấu gạch dưới: `ten_khach`, `so_dien_thoai`
- Không dùng khoảng trắng, không bắt đầu bằng số
- Tên phải có nghĩa: `x` tệ hơn `tuoi_khach`'''))

cells.append(code('''# Gán giá trị vào biến
ten = "Minh"
tuoi = 25
luong = 15000000

# In giá trị của biến
print(ten)
print(tuoi)
print(luong)'''))

cells.append(code('''# Biến có thể thay đổi giá trị
x = 5
print(x)   # In ra: 5

x = 10     # Gán lại giá trị mới
print(x)   # In ra: 10'''))

cells.append(md('''## Phần 5 — Kiểu dữ liệu: 4 loại cơ bản

| Loại | Tên gọi | Ví dụ trong code | Dùng khi |
|---|---|---|---|
| Số nguyên | `int` | `25`, `100`, `-3` | Tuổi, số lượng, năm |
| Số thập phân | `float` | `9.5`, `3.14` | Điểm, giá tiền, cân nặng |
| Chữ/Chuỗi | `str` | `"Minh"`, `"Hà Nội"` | Tên, địa chỉ, tên file |
| Đúng/Sai | `bool` | `True`, `False` | Trạng thái, có/không |

> **Quan trọng:** `500` (số) khác `"500"` (chữ).
> Nếu cần tính toán dùng số. Nếu là tên file hay tên cột dùng chữ trong dấu `" "`.'''))

cells.append(code('''# Các kiểu dữ liệu cơ bản
so_luong = 5                     # int
diem_thi = 8.5                   # float
ho_ten = "Nguyễn Văn Minh"      # str
da_thanh_toan = True             # bool

# Kiểm tra kiểu dữ liệu bằng type()
print(type(so_luong))       # <class 'int'>
print(type(diem_thi))       # <class 'float'>
print(type(ho_ten))         # <class 'str'>
print(type(da_thanh_toan))  # <class 'bool'>'''))

cells.append(code('''# Phép tính với số
a = 10
b = 3

print(a + b)   # Cộng         → 13
print(a - b)   # Trừ          → 7
print(a * b)   # Nhân         → 30
print(a / b)   # Chia         → 3.333...
print(a // b)  # Phần nguyên  → 3
print(a % b)   # Phần dư      → 1
print(a ** b)  # Lũy thừa 10³ → 1000'''))

cells.append(code('''# In chữ kết hợp biến — dùng f-string
ten = "Minh"
tuoi = 25
print(f"Tên tôi là {ten}, tôi {tuoi} tuổi")

# Ví dụ thực tế: tính hóa đơn
don_gia = 50000
so_luong = 3
tong_tien = don_gia * so_luong
print(f"Tổng tiền: {tong_tien:,} đồng")'''))

cells.append(md('''## Phần 6 — Cách đặt yêu cầu với Claude để được code tốt

Đây là kỹ năng **quan trọng nhất** của buổi 1.

**Yêu cầu mơ hồ → code sai hoặc không dùng được:**
> "Viết code xử lý Excel cho tôi"

**Yêu cầu tốt → code dùng được ngay:**
> "Viết code Python xử lý file Excel tên `doanh_thu.xlsx` nằm ở Desktop. File có các cột: A=Ngày, B=Tên khách, C=Sản phẩm, D=Số lượng, E=Đơn giá. Tính thêm cột F=Thành tiền (D×E), lọc các dòng có Thành tiền > 5.000.000, xuất kết quả ra file mới tên `loc_doanh_thu.xlsx`."

---

**Template đặt yêu cầu — copy và điền vào:**

```
Viết code Python để [làm gì].

Input: [tên file / loại file / file ở đâu]
Cấu trúc dữ liệu: [các cột, các trường thông tin]
Output mong muốn: [kết quả trông như thế nào]
Điều kiện đặc biệt: [lọc, sắp xếp, bỏ qua gì...]

Thêm ghi chú tiếng Việt vào code giải thích từng bước.
```'''))

cells.append(md('''## Bài tập Buổi 1

**Bài 1:** Sửa các biến trong ô code bên dưới cho đúng với thông tin của bạn, rồi chạy.

**Bài 2:** Nhờ Claude viết code tính lương cho 3 nhân viên với mức lương khác nhau. Paste code vào ô mới và chạy thử.

**Bài 3:** Dùng template ở Phần 6, đặt yêu cầu với Claude để xử lý 1 file Excel thật trong công việc của bạn.'''))

cells.append(code('''# ← SỬA CÁC BIẾN NÀY cho đúng với thông tin của bạn
ho_ten = "Nguyễn Văn Minh"
phong_ban = "Kinh doanh"
luong_co_ban = 12000000
phu_cap = 2500000

# Phần tính toán — không cần sửa
luong_thuc_nhan = luong_co_ban + phu_cap
thue_tncn = luong_thuc_nhan * 0.1
luong_sau_thue = luong_thuc_nhan - thue_tncn

print(f"Nhân viên: {ho_ten}")
print(f"Phòng ban: {phong_ban}")
print(f"Lương thực nhận: {luong_thuc_nhan:,} đồng")
print(f"Thuế TNCN (10%): {thue_tncn:,} đồng")
print(f"Nhận tay: {luong_sau_thue:,} đồng")'''))

# ── BUỔI 2 ─────────────────────────────────────────────────────────────────
cells.append(md('''---
---
# BUỔI 2 — Đọc Hiểu Code và Xử Lý Khi Có Lỗi

**Mục tiêu buổi này:**
- Đọc hiểu code Claude viết (không cần hiểu từng dòng, hiểu được luồng chính)
- Biết sửa những phần cần tùy chỉnh
- Đọc và xử lý thông báo lỗi'''))

cells.append(md('''## Phần 1 — Đọc code như đọc công thức nấu ăn

Code chạy từ trên xuống. Khi đọc, hỏi 3 câu:
1. **Dữ liệu vào là gì?** (Đọc file gì, lấy thông tin gì)
2. **Làm gì với dữ liệu?** (Tính toán, lọc, sắp xếp)
3. **Kết quả ra là gì?** (Lưu file, in ra màn hình)

Ô code bên dưới là ví dụ minh họa — đọc để hiểu cấu trúc 3 khối, không cần chạy.'''))

cells.append(code('''# ══════════════════════════════════════════════════════
# VÍ DỤ MINH HỌA — cần file Excel thật để chạy được
# Mục đích: học cách ĐỌC cấu trúc code 3 khối
# ══════════════════════════════════════════════════════

# import pandas as pd

# ① ĐỌC FILE VÀO
# du_lieu = pd.read_excel("doanh_thu.xlsx")

# ② XỬ LÝ DỮ LIỆU
# du_lieu["Thanh_tien"] = du_lieu["So_luong"] * du_lieu["Don_gia"]
# ket_qua = du_lieu[du_lieu["Thanh_tien"] > 5000000]

# ③ XUẤT KẾT QUẢ
# ket_qua.to_excel("loc_doanh_thu.xlsx", index=False)
# print(f"Xong! Tìm thấy {len(ket_qua)} dòng.")

print("Cấu trúc 3 khối: ① Đọc vào → ② Xử lý → ③ Xuất ra")'''))

cells.append(md('''## Phần 2 — Điều kiện trong code: if / elif / else

Khi Claude viết code có logic kiểm tra, bạn sẽ thấy `if/else`.

Đọc như văn nói:
- `if (điều kiện):` → "Nếu điều kiện đúng thì..."
- `elif (điều kiện khác):` → "Nếu không, kiểm tra tiếp..."
- `else:` → "Còn lại thì..."

Máy tính kiểm tra **từ trên xuống**, gặp điều kiện đúng đầu tiên thì dừng.

**Khi nào bạn cần sửa?** Khi ngưỡng số liệu không đúng với thực tế — chỉ đổi con số, không đụng gì khác.'''))

cells.append(code('''# Thử đổi giá trị doanh_thu để xem kết quả thay đổi
doanh_thu = 75000000  # ← đổi số này và chạy lại

if doanh_thu > 100000000:
    xep_loai = "Xuất sắc"
elif doanh_thu > 50000000:
    xep_loai = "Tốt"
else:
    xep_loai = "Cần cải thiện"

print(f"Doanh thu: {doanh_thu:,} đồng → Xếp loại: {xep_loai}")'''))

cells.append(code('''# Ví dụ thực tế: phân loại đơn hàng, tính chiết khấu
gia_tri_don_hang = 3500000  # ← đổi số này và chạy lại

if gia_tri_don_hang >= 10000000:
    muc_uu_tien = "Cao — xử lý trong ngày"
    pct_ck = 10
elif gia_tri_don_hang >= 5000000:
    muc_uu_tien = "Trung bình — xử lý trong 2 ngày"
    pct_ck = 5
else:
    muc_uu_tien = "Thường — xử lý trong 3 ngày"
    pct_ck = 0

chiet_khau = gia_tri_don_hang * pct_ck / 100
thanh_toan = gia_tri_don_hang - chiet_khau

print(f"Đơn hàng  : {gia_tri_don_hang:,} đồng")
print(f"Ưu tiên   : {muc_uu_tien}")
print(f"Chiết khấu: {pct_ck}% = {chiet_khau:,} đồng")
print(f"Thanh toán: {thanh_toan:,} đồng")'''))

cells.append(md('''## Phần 3 — Vòng lặp trong code: for

Khi Claude xử lý nhiều dòng dữ liệu hoặc nhiều file, sẽ có vòng lặp.

Đọc như văn nói: *"Với mỗi item trong danh sách, làm những việc bên dưới."*

Vòng `for` = tự động lặp lại, không cần làm tay.'''))

cells.append(code('''# Xử lý từng nhân viên trong danh sách
nhan_vien = [
    {"ten": "Minh", "luong": 15000000},
    {"ten": "An",   "luong": 12000000},
    {"ten": "Binh", "luong": 18000000},
]

print("BANG LUONG THANG 5/2024")
print("-" * 50)

for nv in nhan_vien:
    thue = nv["luong"] * 0.1
    nhan_tay = nv["luong"] - thue
    print(f"{nv['ten']}: Luong {nv['luong']:,} → Nhan tay {nhan_tay:,}")'''))

cells.append(code('''# Mô phỏng: xử lý nhiều file (không cần file thật)
danh_sach_file = [
    "bao_cao_HN.xlsx",
    "bao_cao_HCM.xlsx",
    "bao_cao_DN.xlsx",
]

for ten_file in danh_sach_file:
    print(f"✓ Đang xử lý: {ten_file}")

print(f"\nHoàn thành! Đã xử lý {len(danh_sach_file)} file.")'''))

cells.append(md('''## Phần 4 — Đọc và xử lý thông báo lỗi

Khi chạy code bị lỗi, màn hình hiện thông báo. **Đừng hoảng** — đây là thông tin để tìm vấn đề.

**Cấu trúc thông báo lỗi:**
```
Traceback (most recent call last):
  File "code.py", line 5        ← lỗi ở dòng nào
FileNotFoundError: No such file ← lỗi là gì
```

**3 lỗi phổ biến nhất:**

| Thông báo lỗi | Nghĩa là | Cách xử lý |
|---|---|---|
| `FileNotFoundError` | Không tìm thấy file | Kiểm tra tên file, đường dẫn |
| `ModuleNotFoundError` | Thiếu thư viện | Chạy: `pip install tên_thư_viện` |
| `KeyError: "Tên_cột"` | Không tìm thấy cột | Kiểm tra tên cột trong Excel có khớp không |

**Quy trình xử lý lỗi với Claude:**
1. Copy toàn bộ thông báo lỗi (chữ đỏ)
2. Paste vào Claude kèm: *"Tôi chạy code bị lỗi này, file Excel có các cột: [liệt kê], để ở [đường dẫn]"*
3. Claude sẽ giải thích và sửa'''))

cells.append(code('''# Thử tạo lỗi để xem thông báo trông như thế nào
# Bỏ dấu # ở dòng dưới → nhấn Run → xem lỗi thật

# open("file_khong_ton_tai.xlsx")

print("Bỏ dấu # ở dòng trên để xem thông báo lỗi thật")
print("Sau khi xem, thêm # lại để code chạy bình thường")'''))

cells.append(md('''## Phần 5 — Cài thư viện cần thiết

Thư viện là bộ công cụ mở rộng Python. Cần cài một lần, dùng mãi.

| Thư viện | Dùng cho |
|---|---|
| `pandas` | Xử lý bảng dữ liệu, Excel, CSV |
| `openpyxl` | Đọc/viết file Excel (.xlsx) |
| `python-docx` | Đọc/viết file Word (.docx) |
| `os`, `shutil` | Quản lý file, thư mục (có sẵn, không cần cài) |

**Chạy ô code dưới đây một lần duy nhất:**'''))

cells.append(code('''# Cài thư viện cần thiết — chỉ chạy 1 lần duy nhất
!pip install pandas openpyxl python-docx'''))

cells.append(md('''## Phần 6 — Những chỗ thường phải tự sửa trong code Claude viết

Claude không biết máy tính của bạn, nên để placeholder. Bạn phải tự sửa phần cấu hình:

```python
# ← NHỮNG CHỖ BẠN CẦN SỬA (thường ở đầu file)
ten_file_input = "ten_file_cua_ban.xlsx"      # tên file thật
duong_dan      = "C:/Users/TenCuaBan/Desktop/" # đường dẫn thật
ten_cot_ngay   = "Ngày"                        # đúng với tên cột trong file
ten_file_output = "ket_qua.xlsx"               # tên file kết quả bạn muốn
```

**Cách tìm đường dẫn trên Windows:**
Mở thư mục → nhấp thanh địa chỉ trên cùng → copy.

**Cách tìm tên cột chính xác:**
Mở Excel → nhìn hàng đầu tiên → copy đúng tên (kể cả dấu cách, viết hoa/thường).'''))

cells.append(code('''# Thực hành: nhận dữ liệu từ người dùng
# input() dừng chương trình và chờ bạn gõ vào

ten = input("Bạn tên gì? ")
phong_ban = input("Phòng ban của bạn? ")

print(f"\nXin chào {ten} từ phòng {phong_ban}!")
print("Code đã nhận thông tin của bạn thành công.")'''))

cells.append(md('''## Bài tập Buổi 2

1. **Sửa ngưỡng:** Đổi con số trong ô if/else (Phần 2) cho phù hợp tiêu chuẩn công ty bạn. Chạy lại.
2. **Thêm nhân viên:** Thêm 1 người vào danh sách ở Phần 3. Chạy lại, kiểm tra kết quả.
3. **Xem lỗi thật:** Bỏ dấu `#` ở ô Phần 4. Chạy, đọc thông báo, xác định loại lỗi.
4. **Nhờ Claude:** *"Viết code Python đọc file Excel và in ra: tổng số dòng, tên các cột, 5 dòng đầu tiên. File tên [tên file thật], để ở [đường dẫn thật]."* Paste và chạy thử.'''))

# ── BUỔI 3 ─────────────────────────────────────────────────────────────────
cells.append(md('''---
---
# BUỔI 3 — Thực Hành Tác Vụ Văn Phòng Thật Sự

**Mục tiêu buổi này:**
- Làm được ít nhất 2 tác vụ tự động hóa thật sự hữu ích
- Có quy trình làm việc với Claude ổn định
- Biết mở rộng sang các tác vụ tương tự'''))

cells.append(md('''## Phần 1 — Quy trình làm việc chuẩn với Claude

```
1. XÁC ĐỊNH vấn đề cụ thể (làm tay mất bao lâu? làm gì?)
         ↓
2. MÔ TẢ cho Claude (dùng template buổi 1)
         ↓
3. NHẬN code từ Claude, đọc lướt 3 khối (vào / xử lý / ra)
         ↓
4. SỬA các biến: đường dẫn, tên file, tên cột
         ↓
5. CHẠY THỬ với 5–10 dòng dữ liệu mẫu trước
         ↓
6. KIỂM TRA kết quả có đúng không
         ↓
7. CHẠY với toàn bộ dữ liệu thật
         ↓
8. NẾU LỖI → paste lỗi vào Claude → lặp lại từ bước 3
```'''))

cells.append(md('''## Phần 2 — Tác vụ 1: Gộp nhiều file Excel thành 1

**Tình huống thực tế:** Mỗi tuần nhận báo cáo từ 5 chi nhánh, mỗi chi nhánh 1 file. Cần gộp thành 1 file tổng hợp.

**Cách nhờ Claude:**
> *"Viết code Python gộp tất cả file Excel trong thư mục `C:/BaoCao/Tuan1/` thành 1 file. Thêm cột Nguon ghi tên file gốc vào mỗi dòng. Lưu kết quả ra `tong_hop.xlsx` cùng thư mục."*

Đọc code bên dưới theo 3 khối: **Cấu hình → Đọc & Gộp → Xuất**'''))

cells.append(code('''import pandas as pd
import os

# ══════════════════════════════════════════
# CẤU HÌNH — CHỈ SỬA PHẦN NÀY
# ══════════════════════════════════════════
thu_muc = "C:/BaoCao/Tuan1/"   # ← đổi thành đường dẫn thật
file_output = "tong_hop.xlsx"

# ══════════════════════════════════════════
# ① ĐỌC VÀ GỘP TẤT CẢ FILE
# ══════════════════════════════════════════
tat_ca_du_lieu = []

for ten_file in os.listdir(thu_muc):
    if ten_file.endswith(".xlsx") and ten_file != file_output:
        duong_dan = os.path.join(thu_muc, ten_file)
        df = pd.read_excel(duong_dan)
        df["Nguon"] = ten_file            # thêm cột ghi nguồn gốc
        tat_ca_du_lieu.append(df)
        print(f"✓ Đọc xong: {ten_file}")

# ══════════════════════════════════════════
# ② XUẤT KẾT QUẢ
# ══════════════════════════════════════════
ket_qua = pd.concat(tat_ca_du_lieu, ignore_index=True)
ket_qua.to_excel(os.path.join(thu_muc, file_output), index=False)

print(f"\n✅ Tổng hợp {len(ket_qua)} dòng từ {len(tat_ca_du_lieu)} file.")
print(f"File kết quả: {thu_muc}{file_output}")'''))

cells.append(md('''## Phần 3 — Tác vụ 2: Lọc và phân tích dữ liệu Excel

**Tình huống thực tế:** File 300 đơn hàng. Cần lọc đơn chưa thanh toán, tính tổng tiền còn nợ.

**Cách nhờ Claude:**
> *"Viết code Python đọc file `don_hang.xlsx`. Cột E là Trạng thái (Đã thanh toán / Chưa thanh toán), cột F là Số tiền. Lọc các dòng chưa thanh toán, tính tổng tiền còn nợ, xuất ra `chua_thanh_toan.xlsx` và in tổng tiền."*'''))

cells.append(code('''import pandas as pd

# ══════════════════════════════════════════
# CẤU HÌNH — CHỈ SỬA PHẦN NÀY
# ══════════════════════════════════════════
file_input       = "don_hang.xlsx"          # ← tên file thật
file_output      = "chua_thanh_toan.xlsx"
cot_trang_thai   = "Trang thai"             # ← đúng tên cột trong file
gia_tri_chua_tt  = "Chua thanh toan"        # ← đúng giá trị trong cột
cot_so_tien      = "So tien"                # ← đúng tên cột trong file

# ① ĐỌC FILE
du_lieu = pd.read_excel(file_input)
print(f"Tổng số đơn hàng: {len(du_lieu)}")

# ② LỌC VÀ TÍNH
chua_tt = du_lieu[du_lieu[cot_trang_thai] == gia_tri_chua_tt]
tong_no = chua_tt[cot_so_tien].sum()

# ③ XUẤT KẾT QUẢ
chua_tt.to_excel(file_output, index=False)
print(f"Số đơn chưa thanh toán: {len(chua_tt)}")
print(f"Tổng tiền còn nợ: {tong_no:,.0f} đồng")
print(f"Đã xuất danh sách ra: {file_output}")'''))

cells.append(md('''## Phần 4 — Tác vụ 3: Đổi tên hàng loạt file

**Tình huống thực tế:** 200 file ảnh hóa đơn tên lộn xộn. Cần đổi theo format chuẩn `HD_YYYY-MM-DD_001.jpg`

**Cách nhờ Claude:**
> *"Viết code Python đổi tên tất cả file .jpg trong `C:/HoaDon/`. Trước khi đổi, in danh sách tên cũ → tên mới để tôi kiểm tra."*

> ⚠️ **Luôn yêu cầu Claude in danh sách trước khi thực sự đổi tên.** Kiểm tra đúng ý rồi mới xác nhận.'''))

cells.append(code('''import os
from datetime import datetime

# ══════════════════════════════════════════
# CẤU HÌNH — CHỈ SỬA PHẦN NÀY
# ══════════════════════════════════════════
thu_muc       = "C:/HoaDon/"   # ← đổi thành đường dẫn thật
dinh_dang     = ".jpg"         # ← đổi nếu cần (.png, .pdf...)

# ① LẤY DANH SÁCH VÀ IN ĐỂ KIỂM TRA TRƯỚC
tat_ca_file = sorted([f for f in os.listdir(thu_muc) if f.endswith(dinh_dang)])

print("DANH SÁCH ĐỔI TÊN — kiểm tra trước khi chạy thật:")
print("-" * 60)

danh_sach_doi = []
for i, ten_cu in enumerate(tat_ca_file, start=1):
    duong_dan = os.path.join(thu_muc, ten_cu)
    ngay_sua = datetime.fromtimestamp(os.path.getmtime(duong_dan))
    ten_moi = f"HD_{ngay_sua.strftime('%Y-%m-%d')}_{i:03d}{dinh_dang}"
    danh_sach_doi.append((ten_cu, ten_moi))
    print(f"  {ten_cu}  →  {ten_moi}")

# ② XÁC NHẬN TRƯỚC KHI ĐỔI THẬT
xac_nhan = input("\nGõ 'co' để xác nhận đổi tên, gõ khác để hủy: ")

if xac_nhan.lower() == "co":
    for ten_cu, ten_moi in danh_sach_doi:
        os.rename(os.path.join(thu_muc, ten_cu), os.path.join(thu_muc, ten_moi))
    print(f"✅ Đã đổi tên {len(danh_sach_doi)} file thành công!")
else:
    print("Đã hủy. Không có file nào bị thay đổi.")'''))

cells.append(md('''## Phần 5 — Tác vụ 4: Tạo báo cáo Word tự động

**Tình huống thực tế:** Mỗi tháng tạo báo cáo Word theo mẫu, chỉ thay số liệu.

**Cách nhờ Claude:**
> *"Viết code Python tạo file Word báo cáo tháng. Tiêu đề Báo Cáo Doanh Thu Tháng [tháng]/[năm], bảng tổng hợp 3 cột (Sản phẩm, Số lượng, Doanh thu), đoạn tổng kết. Dữ liệu từ `thang_5.xlsx`. Lưu ra `BaoCao_T5_2024.docx`."*'''))

cells.append(code('''from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ══════════════════════════════════════════
# CẤU HÌNH — CHỈ SỬA PHẦN NÀY
# ══════════════════════════════════════════
thang       = 5
nam         = 2024
ten_cong_ty = "Cong ty ABC"           # ← tên công ty bạn
file_output = f"BaoCao_T{thang}_{nam}.docx"

# Dữ liệu mẫu — thay bằng đọc từ Excel trong thực tế
du_lieu = [
    {"San_pham": "San pham A", "So_luong": 150, "Doanh_thu": 75000000},
    {"San_pham": "San pham B", "So_luong": 80,  "Doanh_thu": 48000000},
    {"San_pham": "San pham C", "So_luong": 200, "Doanh_thu": 60000000},
]

# TẠO FILE WORD
doc = Document()

tieu_de = doc.add_heading(f"BAO CAO DOANH THU THANG {thang}/{nam}", level=1)
tieu_de.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Don vi: {ten_cong_ty}")
doc.add_paragraph(f"Ngay xuat: {datetime.now().strftime('%d/%m/%Y')}")
doc.add_paragraph("")

doc.add_heading("I. KET QUA KINH DOANH", level=2)
bang = doc.add_table(rows=1, cols=3)
bang.style = "Table Grid"

header = bang.rows[0].cells
header[0].text = "San pham"
header[1].text = "So luong"
header[2].text = "Doanh thu (dong)"

tong = 0
for hang in du_lieu:
    dong = bang.add_row().cells
    dong[0].text = hang["San_pham"]
    dong[1].text = str(hang["So_luong"])
    dong[2].text = f"{hang['Doanh_thu']:,}"
    tong += hang["Doanh_thu"]

doc.add_paragraph("")
doc.add_heading("II. TONG KET", level=2)
doc.add_paragraph(f"Tong doanh thu thang {thang}/{nam}: {tong:,} dong")

doc.save(file_output)
print(f"✅ Da tao bao cao: {file_output}")'''))

cells.append(md('''## Phần 6 — Mở rộng: Các tác vụ có thể tự động hóa

**Quản lý file và thư mục:**
- Tự động sắp xếp file vào thư mục theo ngày tháng
- Tìm và xóa file trùng lặp
- Nén file cũ hơn 30 ngày vào .zip

**Xử lý dữ liệu nâng cao:**
- Đối chiếu 2 file Excel tìm ra sự khác biệt
- Tách 1 file lớn thành nhiều file nhỏ theo điều kiện
- Tự động tô màu ô theo giá trị

**Báo cáo và tài liệu:**
- Tạo biểu đồ từ dữ liệu Excel
- Xuất báo cáo PDF
- Tạo slide PowerPoint từ dữ liệu

---

**Cách tiếp cận mọi tác vụ mới:**
Mô tả vấn đề cho Claude → Nhận code → Đọc 3 khối → Sửa cấu hình → Chạy thử → Kiểm tra'''))

cells.append(md('''## Bài tập Buổi 3

**Chọn 1 việc bạn đang làm tay ở công ty** và thực hiện theo quy trình 8 bước.

Tiêu chí chọn việc:
- Hiện tại mất 15–60 phút mỗi lần làm
- Lặp đi lặp lại hàng tuần hoặc hàng tháng
- Dữ liệu ở dạng Excel hoặc file

Ô code bên dưới để paste code từ Claude vào và chạy:'''))

cells.append(code('''# PASTE CODE TỪ CLAUDE VÀO ĐÂY
# Nhớ sửa các biến cấu hình (tên file, đường dẫn, tên cột) trước khi chạy

print("Paste code của bạn vào đây và xóa dòng print này")'''))

# ── TỔNG KẾT ───────────────────────────────────────────────────────────────
cells.append(md('''---
---
## Tổng kết 3 Buổi

| Buổi | Học gì | Bạn làm được gì |
|---|---|---|
| 1 | Máy tính hoạt động, biến, kiểu dữ liệu, cách đặt yêu cầu Claude | Chạy được code, sửa được biến cơ bản |
| 2 | Đọc hiểu code, điều kiện, vòng lặp, xử lý lỗi | Tự sửa lỗi đơn giản, hiểu code Claude viết |
| 3 | 4 tác vụ văn phòng thực tế, quy trình làm việc với Claude | Tự động hóa được việc thật trong công ty |

**Sau 3 buổi, bạn không phải lập trình viên — nhưng bạn là người dùng Claude hiệu quả hơn 80% người dùng thông thường.**

---

## Câu hỏi thường gặp

**"Tôi cần học thêm gì sau 3 buổi này?"**
Không cần học lý thuyết thêm. Cứ gặp vấn đề thật → nhờ Claude giải quyết → đó chính là cách học tốt nhất.

**"Nếu tôi không hiểu code Claude viết thì sao?"**
Paste code vào Claude và hỏi: *"Giải thích code này bằng tiếng Việt đơn giản, từng bước."*

**"Code chạy đúng rồi nhưng muốn thêm tính năng thì làm sao?"**
Mô tả thêm cho Claude: *"Code đang làm được X, tôi muốn thêm Y. Sửa lại giúp tôi."*

**"Có phải nhớ các lệnh không?"**
Không. Claude nhớ thay bạn. Bạn chỉ cần nhớ **cái gì có thể làm được** và **cách mô tả yêu cầu rõ ràng.**'''))

# ── WRITE FILE ─────────────────────────────────────────────────────────────
notebook = {
    "nbformat": 4,
    "nbformat_minor": 5,
    "metadata": {
        "kernelspec": {
            "display_name": "Python 3",
            "language": "python",
            "name": "python3"
        },
        "language_info": {
            "codemirror_mode": {"name": "ipython", "version": 3},
            "file_extension": ".py",
            "mimetype": "text/x-python",
            "name": "python",
            "pygments_lexer": "ipython3",
            "version": "3.10.0"
        }
    },
    "cells": cells
}

output_path = r"C:\Users\QUYET THANG\Desktop\claude\hoc_lap_trinh_van_phong.ipynb"
with open(output_path, "w", encoding="utf-8") as f:
    json.dump(notebook, f, ensure_ascii=False, indent=1)

print(f"✅ Notebook saved: {output_path}")
print(f"   Total cells: {len(cells)}")
